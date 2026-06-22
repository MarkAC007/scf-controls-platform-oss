"""
Vendor Report Generation Service (Issue #61, DPSIA Enhancement).

Generates comprehensive 14-section vendor assessment reports in multiple formats:
- Markdown (base format, stored in DB)
- JSON (structured data)
- PDF (via WeasyPrint)
- DOCX (via python-docx)

Reports include risk score breakdowns, claim verifications, CIA controls,
action items, compensating controls, AI analysis, and recommendations.
"""
import io
import logging
import os
from datetime import datetime
from typing import Any, Dict, List, Optional
from uuid import UUID

from sqlalchemy import select, desc
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from models import (
    Vendor, VendorAssessment, VendorCertification,
    VendorReport, VendorResearchResult,
    VendorClaimVerification, VendorCIAControl,
    VendorActionItem, VendorCompensatingControl,
)
from services.vendor_risk_scoring import (
    FACTOR_WEIGHTS, RECOMMENDATIONS, get_risk_level
)

logger = logging.getLogger(__name__)


async def generate_report(
    db: AsyncSession,
    vendor_id: str,
    organization_id: str,
    user_id: Optional[str] = None,
    assessment_id: Optional[str] = None,
    report_type: str = "comprehensive",
) -> VendorReport:
    """
    Generate a comprehensive vendor assessment report.

    Pulls data from vendor, assessments, certifications, research,
    claim verifications, CIA controls, action items, and compensating
    controls to produce a 14-section DPSIA-quality report.
    """
    # Load vendor with relationships
    vendor_result = await db.execute(
        select(Vendor).where(Vendor.id == vendor_id).options(
            selectinload(Vendor.certifications),
        )
    )
    vendor = vendor_result.scalar_one_or_none()
    if not vendor:
        raise ValueError(f"Vendor {vendor_id} not found")

    # Load the specific assessment or the most recent completed one
    if assessment_id:
        assess_result = await db.execute(
            select(VendorAssessment).where(VendorAssessment.id == assessment_id)
        )
        assessment = assess_result.scalar_one_or_none()
    else:
        assess_result = await db.execute(
            select(VendorAssessment)
            .where(VendorAssessment.vendor_id == vendor_id)
            .order_by(desc(VendorAssessment.assessment_date))
            .limit(1)
        )
        assessment = assess_result.scalar_one_or_none()

    # Load latest research results
    research_result = await db.execute(
        select(VendorResearchResult)
        .where(
            VendorResearchResult.vendor_id == vendor_id,
            VendorResearchResult.status.in_(["completed", "partial"]),
        )
        .order_by(desc(VendorResearchResult.created_at))
        .limit(1)
    )
    research = research_result.scalar_one_or_none()

    # Load DPSIA Enhancement data
    verif_result = await db.execute(
        select(VendorClaimVerification)
        .where(VendorClaimVerification.vendor_id == vendor_id)
        .order_by(VendorClaimVerification.created_at.desc())
    )
    claim_verifications = verif_result.scalars().all()

    cia_controls = []
    if assessment:
        cia_result = await db.execute(
            select(VendorCIAControl).where(VendorCIAControl.assessment_id == assessment.id)
        )
        cia_controls = cia_result.scalars().all()

    action_result = await db.execute(
        select(VendorActionItem)
        .where(VendorActionItem.vendor_id == vendor_id)
        .order_by(VendorActionItem.created_at.desc())
    )
    action_items = action_result.scalars().all()

    comp_result = await db.execute(
        select(VendorCompensatingControl)
        .where(VendorCompensatingControl.vendor_id == vendor_id)
    )
    compensating_controls = comp_result.scalars().all()

    # Build report content
    now = datetime.utcnow()
    risk_score = assessment.final_risk_score if assessment else (vendor.risk_score or 0)
    risk_level_val = assessment.risk_level if assessment else (vendor.risk_level or "medium")
    recommendation = RECOMMENDATIONS.get(risk_level_val, "Approve with monitoring")

    markdown = _build_markdown_report(
        vendor=vendor,
        assessment=assessment,
        research=research,
        risk_score=risk_score,
        risk_level=risk_level_val,
        recommendation=recommendation,
        report_date=now,
        claim_verifications=claim_verifications,
        cia_controls=cia_controls,
        action_items=action_items,
        compensating_controls=compensating_controls,
    )

    content_json = _build_json_content(
        vendor=vendor,
        assessment=assessment,
        research=research,
        risk_score=risk_score,
        risk_level=risk_level_val,
        recommendation=recommendation,
        claim_verifications=claim_verifications,
        cia_controls=cia_controls,
        action_items=action_items,
        compensating_controls=compensating_controls,
    )

    # Check for existing reports to set version
    existing_count_result = await db.execute(
        select(VendorReport)
        .where(VendorReport.vendor_id == vendor_id)
    )
    existing_reports = existing_count_result.scalars().all()
    version = len(existing_reports) + 1

    title = f"Vendor Security Assessment Report - {vendor.name} - v{version}"

    # Create report record
    report = VendorReport(
        vendor_id=UUID(vendor_id) if isinstance(vendor_id, str) else vendor_id,
        assessment_id=UUID(assessment_id) if assessment_id else (assessment.id if assessment else None),
        organization_id=UUID(organization_id) if isinstance(organization_id, str) else organization_id,
        report_type=report_type,
        title=title,
        content_markdown=markdown,
        content_json=content_json,
        risk_score=risk_score,
        risk_level=risk_level_val,
        recommendation=recommendation,
        version=version,
        generated_by_user_id=UUID(user_id) if user_id else None,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)

    return report


def _build_markdown_report(
    vendor: Vendor,
    assessment: Optional[VendorAssessment],
    research: Optional[VendorResearchResult],
    risk_score: int,
    risk_level: str,
    recommendation: str,
    report_date: datetime,
    claim_verifications: Optional[list] = None,
    cia_controls: Optional[list] = None,
    action_items: Optional[list] = None,
    compensating_controls: Optional[list] = None,
) -> str:
    """Build the 14-section DPSIA-quality markdown report."""

    rag_map = {"low": "Green", "medium": "Amber", "high": "Red", "critical": "Red"}
    rag_status = rag_map.get(risk_level, "Amber")
    rag_emoji = {"Green": "✅", "Amber": "⚠️", "Red": "❌"}.get(rag_status, "⚠️")

    sections = []

    # ── Section 1: Executive Summary ──
    sections.append("# Vendor Security Assessment Report\n")
    sections.append("## 1. Executive Summary\n")
    sections.append(f"| Field | Value |")
    sections.append(f"|-------|-------|")
    sections.append(f"| **Vendor** | {vendor.name} |")
    sections.append(f"| **Domain** | {vendor.website or 'N/A'} |")
    sections.append(f"| **Assessment Date** | {report_date.strftime('%d %B %Y')} |")
    sections.append(f"| **RAG Status** | {rag_emoji} {rag_status} |")
    sections.append(f"| **Risk Score** | {risk_score}/25 ({risk_level.title()}) |")
    sections.append(f"| **Recommendation** | {recommendation} |")
    sections.append("")

    if assessment and assessment.ai_analysis:
        sections.append(assessment.ai_analysis)
        sections.append("")

    # ── Section 2: Vendor Overview ──
    sections.append("## 2. Vendor Overview\n")
    sections.append(f"| Field | Value |")
    sections.append(f"|-------|-------|")
    sections.append(f"| **Name** | {vendor.name} |")
    sections.append(f"| **Category** | {vendor.category or 'N/A'} |")
    sections.append(f"| **Status** | {vendor.status.replace('_', ' ').title()} |")
    sections.append(f"| **Criticality** | {vendor.criticality.title()} |")
    sections.append(f"| **Data Classification** | {(vendor.data_classification or 'N/A').title()} |")
    sections.append(f"| **Contact** | {vendor.contact_name or 'N/A'} ({vendor.contact_email or 'N/A'}) |")
    if vendor.contract_start_date:
        sections.append(f"| **Contract Period** | {vendor.contract_start_date} to {vendor.contract_end_date or 'Ongoing'} |")
    sections.append("")

    # ── Section 3: Certification Verification ──
    sections.append("## 3. Certification Verification\n")
    if vendor.certifications:
        sections.append("| Certification | Status | Verified | Expiry |")
        sections.append("|--------------|--------|----------|--------|")
        for cert in vendor.certifications:
            verified = "—"
            if claim_verifications:
                matching = [v for v in claim_verifications
                           if v.claim_type == "certification"
                           and cert.certification_name in (v.claim_description or "")]
                if matching:
                    v = matching[0]
                    status_emoji = {"confirmed": "✅", "unverified": "⚠️", "discrepancy": "❌", "anomaly": "❌"}
                    verified = f"{status_emoji.get(v.verification_status, '⚠️')} {v.verification_status.title()}"
            expiry = cert.expiry_date.strftime('%d %b %Y') if cert.expiry_date else 'N/A'
            sections.append(f"| {cert.certification_name} | {cert.status.title()} | {verified} | {expiry} |")
    else:
        sections.append("No certifications on record.")
    sections.append("")

    # ── Section 4: Breach History ──
    sections.append("## 4. Breach History\n")
    if research and research.hibp_results:
        breaches = (research.hibp_results or {}).get("breaches") or []
        if breaches:
            sections.append(f"**{len(breaches)} breach(es) identified** (Total accounts affected: {(research.hibp_results or {}).get('total_pwned_accounts', 'Unknown')})\n")
            sections.append("| Breach | Date | Accounts | Data Classes | Sensitive |")
            sections.append("|--------|------|----------|-------------|-----------|")
            for b in breaches[:15]:
                name = b.get("name") or b.get("Name") or "Unknown"
                date_str = b.get("date") or b.get("BreachDate") or "Unknown"
                pwn_count = b.get("pwn_count") or b.get("PwnCount") or 0
                data_classes = ", ".join((b.get("data_classes") or b.get("DataClasses") or [])[:4])
                sensitive = "Yes" if b.get("is_sensitive") or b.get("IsSensitive") else "No"
                sections.append(f"| {name} | {date_str} | {pwn_count:,} | {data_classes} | {sensitive} |")
        else:
            sections.append("✅ No known breaches found via Have I Been Pwned.")
    else:
        sections.append("No breach data available. Run vendor research to check HIBP.")
    sections.append("")

    # ── Section 5: CVE/Vulnerability History ──
    sections.append("## 5. CVE/Vulnerability History\n")
    if research and research.cve_nvd_results:
        cve_data = research.cve_nvd_results or {}
        cves = cve_data.get("cves") or cve_data.get("vulnerabilities") or []
        severity_counts = cve_data.get("severity_counts") or {}
        if cves:
            sections.append(f"**{cve_data.get('total_results', len(cves))} CVE(s) found** — "
                          f"Critical: {severity_counts.get('CRITICAL', 0)}, "
                          f"High: {severity_counts.get('HIGH', 0)}, "
                          f"Medium: {severity_counts.get('MEDIUM', 0)}, "
                          f"Low: {severity_counts.get('LOW', 0)}\n")
            sections.append("| CVE ID | CVSS | Severity | Description |")
            sections.append("|--------|------|----------|-------------|")
            for v in cves[:15]:
                cve_id = v.get("cve_id") or "Unknown"
                cvss = v.get("cvss_score") or "N/A"
                severity = v.get("severity") or "Unknown"
                desc_text = (v.get("description") or "")[:120]
                sections.append(f"| {cve_id} | {cvss} | {severity} | {desc_text} |")
        else:
            sections.append("✅ No known CVEs found in NVD.")
    else:
        sections.append("No vulnerability data available.")
    sections.append("")

    # ── Section 6: CISA KEV Analysis ──
    sections.append("## 6. CISA KEV Analysis\n")
    if research and research.cisa_kev_results:
        kev = research.cisa_kev_results or {}
        vulns = kev.get("vulnerabilities") or []
        overdue = kev.get("overdue_count") or 0
        if vulns:
            sections.append(f"**{kev.get('total_matches', len(vulns))} known exploited vulnerabilities** "
                          f"({overdue} overdue for remediation)\n")
            if overdue > 0:
                sections.append(f"⚠️ **{overdue} overdue required action(s)** — immediate attention required.\n")
            sections.append("| CVE ID | Product | Date Added | Due Date | Ransomware Use |")
            sections.append("|--------|---------|-----------|----------|----------------|")
            for v in vulns[:10]:
                sections.append(f"| {v.get('cve_id', 'N/A')} | {v.get('product', 'N/A')} | {v.get('date_added', 'N/A')} | {v.get('due_date', 'N/A')} | {v.get('known_ransomware_use', 'Unknown')} |")
        else:
            sections.append("✅ No CISA KEV entries found for this vendor.")
    else:
        sections.append("No CISA KEV data available.")
    sections.append("")

    # ── Section 7: CIA Triad Assessment ──
    sections.append("## 7. CIA Triad Assessment\n")
    if cia_controls:
        for pillar in ["confidentiality", "integrity", "availability"]:
            pillar_controls = [c for c in cia_controls if c.pillar == pillar]
            if pillar_controls:
                avg = sum(c.score or 0 for c in pillar_controls) / len(pillar_controls)
                sections.append(f"### {pillar.title()} (Average: {avg:.1f}/5)\n")
                sections.append("| Control | Score | Detail |")
                sections.append("|---------|-------|--------|")
                for c in pillar_controls:
                    sections.append(f"| {c.control_name} | {c.score or 'N/A'}/5 | {(c.detail or '-')[:100]} |")
                sections.append("")
    elif assessment and any([assessment.confidentiality_score, assessment.integrity_score, assessment.availability_score]):
        sections.append("| Pillar | Score |")
        sections.append("|--------|-------|")
        sections.append(f"| Confidentiality | {assessment.confidentiality_score or 'N/A'}/5 |")
        sections.append(f"| Integrity | {assessment.integrity_score or 'N/A'}/5 |")
        sections.append(f"| Availability | {assessment.availability_score or 'N/A'}/5 |")
    else:
        sections.append("No CIA assessment data available.")
    sections.append("")

    # ── Section 8: Data Handling Assessment ──
    sections.append("## 8. Data Handling Assessment\n")
    sections.append(f"- **Data Classification:** {(vendor.data_classification or 'Not specified').title()}")
    if vendor.data_classification in ('confidential', 'restricted'):
        sections.append("- ⚠️ High-sensitivity data handling — enhanced controls required")
    sections.append("")

    # ── Section 9: Regulatory Compliance ──
    sections.append("## 9. Regulatory Compliance\n")
    if research and research.regulatory_results:
        findings = (research.regulatory_results or {}).get("findings") or \
                   (research.regulatory_results or {}).get("actions") or \
                   (research.regulatory_results or {}).get("enforcement_actions") or []
        if findings:
            sections.append(f"**{len(findings)} regulatory finding(s)**\n")
            for f in findings[:10]:
                severity = f.get("severity", "info").title()
                desc_text = f.get("description") or f.get("title") or str(f)
                source = f.get("source_url") or f.get("type") or ""
                sections.append(f"- **[{severity}]** {desc_text[:200]}")
                if source:
                    sections.append(f"  Source: {source}")
        else:
            sections.append("✅ No regulatory enforcement actions found.")
    else:
        sections.append("No regulatory data available.")
    sections.append("")

    # ── Section 10: Risk Scoring ──
    sections.append("## 10. Risk Scoring\n")
    if assessment:
        inherent = assessment.inherent_risk_score
        effectiveness = assessment.control_effectiveness_pct
        residual = assessment.final_risk_score

        if inherent and effectiveness is not None:
            sections.append("### Risk Waterfall\n")
            sections.append(f"**Inherent Risk:** {inherent}/25 ({assessment.inherent_risk_level or 'N/A'})")
            sections.append(f"→ **Control Effectiveness:** {effectiveness}%")
            sections.append(f"→ **Residual Risk:** {residual}/25 ({assessment.risk_level or 'N/A'})\n")

        sections.append("### Factor Breakdown\n")
        sections.append("| Factor | Score | Weight | Weighted |")
        sections.append("|--------|-------|--------|----------|")
        breach = assessment.breach_score or 0
        cert = assessment.certification_score or 0
        cve = assessment.cve_score or 0
        reg = assessment.regulatory_score or 0
        data = assessment.data_handling_score or 0
        sections.append(f"| Breach History | {breach}/25 | 25% | {breach * 0.25:.1f} |")
        sections.append(f"| Certifications | {cert}/25 | 20% | {cert * 0.20:.1f} |")
        sections.append(f"| CVE Severity | {cve}/25 | 20% | {cve * 0.20:.1f} |")
        sections.append(f"| Regulatory | {reg}/25 | 15% | {reg * 0.15:.1f} |")
        sections.append(f"| Data Handling | {data}/25 | 20% | {data * 0.20:.1f} |")
        sections.append(f"| **Residual Risk** | | | **{risk_score}/25** |")
    else:
        sections.append("No assessment data available for risk scoring.")
    sections.append("")

    # ── Section 11: Compensating Controls ──
    sections.append("## 11. Compensating Controls\n")
    if compensating_controls:
        sections.append("| Gap | Compensating Control | Effectiveness |")
        sections.append("|-----|---------------------|--------------|")
        for cc in compensating_controls:
            eff_emoji = {"full": "✅", "partial": "⚠️", "minimal": "❌"}.get(cc.effectiveness_rating, "⚠️")
            sections.append(f"| {(cc.gap_description or '')[:80]} | {(cc.compensating_control or '')[:80]} | {eff_emoji} {cc.effectiveness_rating.title()} |")
    else:
        sections.append("No compensating controls documented.")
    sections.append("")

    # ── Section 12: Claim Verification Matrix ──
    sections.append("## 12. Claim Verification Matrix\n")
    if claim_verifications:
        status_counts = {}
        for v in claim_verifications:
            s = v.verification_status
            status_counts[s] = status_counts.get(s, 0) + 1
        summary_parts = [f"{count} {status.title()}" for status, count in status_counts.items()]
        sections.append(f"**Summary:** {', '.join(summary_parts)}\n")
        sections.append("| Claim | Type | Status | Source | Detail |")
        sections.append("|-------|------|--------|--------|--------|")
        for v in claim_verifications:
            status_emoji = {"confirmed": "✅", "unverified": "⚠️", "discrepancy": "❌", "anomaly": "❌"}.get(v.verification_status, "⚠️")
            sections.append(
                f"| {(v.claim_description or '')[:60]} | {v.claim_type} | "
                f"{status_emoji} {v.verification_status.title()} | "
                f"{v.verification_source or '-'} | {(v.verification_detail or '-')[:60]} |"
            )
    else:
        sections.append("No claim verifications performed. Run vendor verification to cross-reference claims.")
    sections.append("")

    # ── Section 13: Action Items ──
    sections.append("## 13. Action Items\n")
    if action_items:
        open_items = [a for a in action_items if a.status in ('open', 'in_progress')]
        sections.append(f"**{len(open_items)} open item(s)** out of {len(action_items)} total\n")
        sections.append("| Title | Priority | Owner | Due Date | Status |")
        sections.append("|-------|----------|-------|----------|--------|")
        priority_emoji = {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🟢"}
        for a in action_items:
            emoji = priority_emoji.get(a.priority, "⚪")
            due = a.due_date.strftime('%d %b %Y') if a.due_date else '-'
            auto_badge = " [Auto]" if a.auto_generated else ""
            sections.append(
                f"| {a.title}{auto_badge} | {emoji} {a.priority.title()} | "
                f"{a.owner_name or '-'} | {due} | {a.status.replace('_', ' ').title()} |"
            )
    else:
        sections.append("No action items recorded.")
    sections.append("")

    # ── Section 14: Appendices ──
    sections.append("## 14. Appendices\n")
    sections.append("### Methodology")
    sections.append("This report was generated using the SCF Controls Platform DPSIA-quality assessment methodology:")
    sections.append("- **Research sources:** HIBP, CISA KEV, CVE/NVD, Regulatory (OFAC, SEC)")
    sections.append("- **Risk model:** Inherent Risk → Control Effectiveness → Residual Risk")
    sections.append("- **Scoring:** Weighted factor model (breach 25%, cert 20%, CVE 20%, regulatory 15%, data handling 20%)")
    sections.append("- **Verification:** Independent cross-referencing of vendor claims against research data\n")
    sections.append("### Data Sources")
    if research:
        sections.append(f"- Research job: {research.job_id}")
        sections.append(f"- Domain researched: {research.researched_domain or 'N/A'}")
        completed = research.completed_at.strftime('%d %B %Y %H:%M UTC') if research.completed_at else 'N/A'
        sections.append(f"- Completed: {completed}")
    sections.append("")
    sections.append("### Limitations")
    sections.append("- Research data reflects point-in-time findings and may not capture recent changes")
    sections.append("- HIBP coverage depends on publicly reported breaches")
    sections.append("- Regulatory findings are limited to sources with public APIs")
    sections.append("")

    # Footer
    sections.append("---")
    sections.append(f"Generated by SCF Controls Platform | {report_date.strftime('%d %B %Y %H:%M UTC')}")
    if assessment:
        sections.append(f"Assessment ID: {assessment.id}")

    return "\n".join(sections)


def _build_json_content(
    vendor: Vendor,
    assessment: Optional[VendorAssessment],
    research: Optional[VendorResearchResult],
    risk_score: int,
    risk_level: str,
    recommendation: str,
    claim_verifications: Optional[list] = None,
    cia_controls: Optional[list] = None,
    action_items: Optional[list] = None,
    compensating_controls: Optional[list] = None,
) -> dict:
    """Build structured JSON content for the report."""
    result = {
        "vendor": {
            "id": str(vendor.id),
            "name": vendor.name,
            "website": vendor.website,
            "category": vendor.category,
            "status": vendor.status,
            "criticality": vendor.criticality,
        },
        "risk_summary": {
            "score": risk_score,
            "level": risk_level,
            "recommendation": recommendation,
            "inherent_risk_score": assessment.inherent_risk_score if assessment else None,
            "inherent_risk_level": assessment.inherent_risk_level if assessment else None,
            "control_effectiveness_pct": assessment.control_effectiveness_pct if assessment else None,
        },
        "factor_scores": {
            "breach_score": assessment.breach_score if assessment else None,
            "certification_score": assessment.certification_score if assessment else None,
            "cve_score": assessment.cve_score if assessment else None,
            "regulatory_score": assessment.regulatory_score if assessment else None,
            "data_handling_score": assessment.data_handling_score if assessment else None,
        } if assessment else None,
        "cia_scores": {
            "confidentiality": assessment.confidentiality_score if assessment else None,
            "integrity": assessment.integrity_score if assessment else None,
            "availability": assessment.availability_score if assessment else None,
        } if assessment else None,
        "certifications": [
            {
                "name": c.certification_name,
                "status": c.status,
                "expiry_date": c.expiry_date.isoformat() if c.expiry_date else None,
            }
            for c in (vendor.certifications or [])
        ],
        "claim_verifications": [
            {
                "claim_type": v.claim_type,
                "claim_description": v.claim_description,
                "verification_status": v.verification_status,
                "verification_source": v.verification_source,
            }
            for v in (claim_verifications or [])
        ],
        "cia_controls": [
            {
                "pillar": c.pillar,
                "control_name": c.control_name,
                "score": c.score,
                "detail": c.detail,
            }
            for c in (cia_controls or [])
        ],
        "action_items": [
            {
                "title": a.title,
                "priority": a.priority,
                "status": a.status,
                "owner_name": a.owner_name,
                "due_date": a.due_date.isoformat() if a.due_date else None,
                "auto_generated": a.auto_generated,
            }
            for a in (action_items or [])
        ],
        "compensating_controls": [
            {
                "gap_description": cc.gap_description,
                "compensating_control": cc.compensating_control,
                "effectiveness_rating": cc.effectiveness_rating,
            }
            for cc in (compensating_controls or [])
        ],
    }
    return result


def export_as_pdf(markdown_content: str, title: str) -> bytes:
    """
    Export report as PDF using WeasyPrint.

    Falls back to a basic HTML-to-PDF if WeasyPrint is unavailable.
    """
    try:
        from weasyprint import HTML
        import markdown

        html_content = markdown.markdown(
            markdown_content,
            extensions=['tables', 'fenced_code']
        )

        styled_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, Helvetica, sans-serif; margin: 40px; color: #333; line-height: 1.6; }}
        h1 {{ color: #1a365d; border-bottom: 2px solid #1a365d; padding-bottom: 10px; }}
        h2 {{ color: #2d4a7a; margin-top: 30px; }}
        h3 {{ color: #4a6fa5; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background-color: #f0f4f8; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #f9fafb; }}
        strong {{ color: #1a365d; }}
        hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 30px 0; }}
        ul, ol {{ padding-left: 20px; }}
        li {{ margin-bottom: 5px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

        pdf_bytes = HTML(string=styled_html).write_pdf()
        return pdf_bytes

    except ImportError as e:
        logger.warning(f"WeasyPrint or markdown not available: {e}. PDF export unavailable.")
        raise ImportError(
            "PDF export requires 'weasyprint' and 'markdown' packages. "
            "Install with: pip install weasyprint markdown"
        )


def export_as_docx(markdown_content: str, title: str) -> bytes:
    """
    Export report as DOCX using python-docx.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()

        # Styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Parse markdown sections
        lines = markdown_content.split('\n')
        in_table = False
        table_rows = []

        for line in lines:
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                if in_table and table_rows:
                    _add_table_to_doc(doc, table_rows)
                    table_rows = []
                    in_table = False
                continue

            # Table rows
            if stripped.startswith('|') and stripped.endswith('|'):
                # Skip separator rows
                if all(c in '|-: ' for c in stripped):
                    continue
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                table_rows.append(cells)
                in_table = True
                continue

            # End table if we were in one
            if in_table and table_rows:
                _add_table_to_doc(doc, table_rows)
                table_rows = []
                in_table = False

            # Headers
            if stripped.startswith('# '):
                heading = doc.add_heading(stripped[2:], level=1)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
            elif stripped.startswith('## '):
                heading = doc.add_heading(stripped[3:], level=2)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0x2d, 0x4a, 0x7a)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('---'):
                doc.add_paragraph('_' * 50)
            elif stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            else:
                # Handle bold markers
                text = stripped.replace('**', '')
                doc.add_paragraph(text)

        # Handle remaining table
        if in_table and table_rows:
            _add_table_to_doc(doc, table_rows)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError as e:
        logger.warning(f"python-docx not available: {e}. DOCX export unavailable.")
        raise ImportError(
            "DOCX export requires 'python-docx' package. "
            "Install with: pip install python-docx"
        )


def _add_table_to_doc(doc, rows):
    """Add a table to a docx document from parsed markdown rows."""
    if not rows:
        return

    try:
        from docx.shared import RGBColor  # noqa: F811
    except ImportError:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text.replace('**', '').strip()
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


async def send_report_email(
    to_email: str,
    to_name: Optional[str],
    report: VendorReport,
    vendor_name: str,
) -> Optional[str]:
    """
    Send a vendor assessment report via email using Resend.

    Returns email ID if successful, None otherwise.
    """
    from services.email_service import RESEND_ENABLED, RESEND_FROM_EMAIL, APP_URL

    if not RESEND_ENABLED:
        logger.warning("Resend not enabled - cannot send report email")
        return None

    try:
        import resend

        rag_map = {"low": "#22c55e", "medium": "#f59e0b", "high": "#ef4444", "critical": "#dc2626"}
        risk_colour = rag_map.get(report.risk_level or "medium", "#f59e0b")

        html_body = f"""
        <html>
        <body style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
            <h1 style="color: #1a365d;">Vendor Assessment Report</h1>
            <p>Hi {to_name or 'there'},</p>
            <p>Please find the vendor assessment report for <strong>{vendor_name}</strong>.</p>

            <div style="background: #f8fafc; border-radius: 8px; padding: 20px; margin: 20px 0; border-left: 4px solid {risk_colour};">
                <p style="margin: 0 0 8px 0;"><strong>Risk Score:</strong> {report.risk_score or 'N/A'}/25</p>
                <p style="margin: 0 0 8px 0;"><strong>Risk Level:</strong> <span style="color: {risk_colour}; font-weight: bold;">{(report.risk_level or 'N/A').title()}</span></p>
                <p style="margin: 0;"><strong>Recommendation:</strong> {report.recommendation or 'N/A'}</p>
            </div>

            <p>View the full report in the platform:</p>
            <p style="margin: 20px 0;">
                <a href="{APP_URL}" style="display: inline-block; background-color: #1a365d; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px;">
                    View Full Report
                </a>
            </p>

            <hr style="border: none; border-top: 1px solid #e2e8f0; margin: 30px 0;">
            <p style="font-size: 13px; color: #888;">Generated by SCF Controls Platform</p>
        </body>
        </html>
        """

        import re
        safe_vendor = re.sub(r'[^a-zA-Z0-9_-]', '_', vendor_name)[:50]

        params = {
            "from": RESEND_FROM_EMAIL,
            "to": [to_email],
            "subject": f"Vendor Assessment Report: {vendor_name}",
            "html": html_body,
            "tags": [
                {"name": "type", "value": "vendor_report"},
                {"name": "vendor", "value": safe_vendor},
            ]
        }

        email = resend.Emails.send(params)
        logger.info(f"Report email sent to {to_email}: {email.get('id', 'unknown')}")
        return email.get('id')

    except Exception as exc:
        logger.error(f"Failed to send report email: {exc}")
        return None
